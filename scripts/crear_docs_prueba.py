"""
Script para generar documentos de prueba realistas.
Crea una resolución sancionatoria base y varios recursos de reposición.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

CARPETA = os.path.join(os.path.dirname(__file__), "prueba")
CARPETA_BASE = CARPETA
CARPETA_RECURSOS = os.path.join(CARPETA, "recursos")
os.makedirs(CARPETA_BASE, exist_ok=True)
os.makedirs(CARPETA_RECURSOS, exist_ok=True)


def agregar_parrafo(doc, texto, negrita=False, centrado=False, tamaño=11):
    p = doc.add_paragraph()
    if centrado:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = negrita
    run.font.size = Pt(tamaño)
    return p


# ─────────────────────────────────────────────────────────────────────────────
# RESOLUCIÓN SANCIONATORIA BASE
# ─────────────────────────────────────────────────────────────────────────────
def crear_resolucion_base():
    doc = Document()

    agregar_parrafo(doc, "RESOLUCIÓN No. 2024-001", negrita=True, centrado=True, tamaño=14)
    agregar_parrafo(doc, "Por la cual se impone sanción administrativa", centrado=True)
    agregar_parrafo(doc, "Bogotá D.C., 15 de enero de 2024", centrado=True)
    doc.add_paragraph()

    agregar_parrafo(doc, "I. ANTECEDENTES", negrita=True)
    agregar_parrafo(doc,
        "La Superintendencia inició investigación administrativa contra la empresa XYZ S.A.S., "
        "identificada con NIT 900.123.456-7, por presunta infracción a las normas del sector. "
        "El proceso se adelantó conforme al procedimiento establecido en la Ley 1437 de 2011.")

    agregar_parrafo(doc, "II. CONSIDERACIONES", negrita=True)

    agregar_parrafo(doc,
        "Con respecto al argumento del debido proceso, esta Entidad considera que el procedimiento "
        "administrativo sancionatorio fue adelantado en estricto cumplimiento de las garantías "
        "constitucionales y legales. Se notificó oportunamente al investigado, se le otorgó traslado "
        "de cargos y se garantizó el derecho de contradicción y defensa.")

    agregar_parrafo(doc,
        "En relación con el argumento de prescripción de la acción sancionatoria, esta Superintendencia "
        "advierte que los hechos objeto de investigación ocurrieron el 10 de marzo de 2022, y la "
        "apertura de investigación se realizó el 5 de septiembre de 2022, es decir, dentro del término "
        "de tres (3) años previsto en el artículo 52 del Código Disciplinario. Por tanto, la acción "
        "sancionatoria no está prescrita.")

    agregar_parrafo(doc,
        "Frente al argumento de falta de competencia, esta Entidad tiene plenas facultades "
        "constitucionales y legales para adelantar investigaciones y sancionar conductas que afecten "
        "el orden económico y social, de conformidad con lo establecido en el Decreto 4170 de 2011 "
        "y la Ley 1480 de 2011.")

    agregar_parrafo(doc,
        "En cuanto al argumento de ausencia de dolo o culpa, esta Superintendencia verificó que la "
        "conducta infractora fue cometida de manera negligente, pues la empresa contaba con los "
        "medios para cumplir con la obligación y omitió adoptar las medidas necesarias. La culpa "
        "leve es suficiente para imponer sanción administrativa.")

    agregar_parrafo(doc, "III. DECISIÓN", negrita=True)
    agregar_parrafo(doc,
        "Por lo anteriormente expuesto, esta Superintendencia RESUELVE: PRIMERO: Sancionar a "
        "XYZ S.A.S. con multa de cien millones de pesos ($100.000.000). SEGUNDO: Notificar la "
        "presente resolución conforme a la ley. TERCERO: Contra la presente resolución proceden "
        "los recursos de reposición y en subsidio de apelación.")

    ruta = os.path.join(CARPETA_BASE, "resolucion_base.docx")
    doc.save(ruta)
    print(f"Creada: {ruta}")


# ─────────────────────────────────────────────────────────────────────────────
# RECURSOS DE REPOSICIÓN
# ─────────────────────────────────────────────────────────────────────────────
recursos = [
    {
        "nombre": "recurso_empresa_xyz.docx",
        "titulo": "RECURSO DE REPOSICIÓN — XYZ S.A.S.",
        "argumentos": [
            ("VIOLACIÓN AL DEBIDO PROCESO",
             "Señala el recurrente que la resolución sancionatoria desconoce las garantías del "
             "debido proceso administrativo, toda vez que la entidad no cumplió con el traslado "
             "oportuno de la totalidad del expediente. Esta omisión vulnera el artículo 29 de la "
             "Constitución Política y el artículo 47 de la Ley 1437 de 2011. En consecuencia, "
             "solicita se declare la nulidad de la actuación."),

            ("PRESCRIPCIÓN DE LA ACCIÓN SANCIONATORIA",
             "Aduce el apoderado que la acción sancionatoria se encuentra prescrita, pues los "
             "hechos investigados ocurrieron el 10 de marzo de 2022 y la notificación del auto "
             "de apertura no se realizó dentro del término de un año establecido en la norma. "
             "Por lo anterior, la potestad sancionatoria se extinguió y la resolución carece "
             "de fundamento jurídico."),

            ("AUSENCIA DE CULPA",
             "Manifiesta el recurrente que no existió culpa ni dolo en la conducta de su "
             "representada, pues adoptó todas las medidas razonables para dar cumplimiento a "
             "las obligaciones legales. La empresa contratió una firma consultora especializada "
             "y actuó de buena fe, lo que excluye cualquier reproche sancionatorio."),

            ("DESPROPORCIÓN DE LA SANCIÓN",
             "Alega el recurrente que la multa impuesta de cien millones de pesos resulta "
             "desproporcionada en relación con la gravedad de la conducta y la capacidad "
             "económica de la empresa. Solicita que se reduzca la sanción a veinte millones "
             "de pesos, atendiendo los criterios de graduación previstos en la ley."),
        ]
    },
    {
        "nombre": "recurso_accionista_minoritario.docx",
        "titulo": "RECURSO DE REPOSICIÓN — CARLOS PÉREZ (Accionista)",
        "argumentos": [
            ("FALTA DE COMPETENCIA DE LA ENTIDAD",
             "El recurrente sostiene que la Superintendencia carece de competencia para "
             "investigar y sancionar las conductas objeto del procedimiento, pues corresponde "
             "exclusivamente a la Contraloría General de la República adelantar este tipo de "
             "actuaciones. La resolución impugnada viola el principio de legalidad al usurpar "
             "funciones que no le corresponden a esta Entidad."),

            ("VIOLACIÓN AL DEBIDO PROCESO Y DERECHO DE DEFENSA",
             "El recurrente manifiesta que durante el trámite administrativo no se le permitió "
             "ejercer plenamente su derecho de contradicción, pues varios documentos del "
             "expediente fueron clasificados como reservados sin justificación legal. Esta "
             "actuación vulnera el artículo 29 constitucional y el principio de publicidad "
             "del procedimiento administrativo."),

            ("PRESCRIPCIÓN DE LA ACCIÓN",
             "Argumenta el recurrente que transcurrió más de tres años entre la ocurrencia "
             "de los hechos y la iniciación formal de la investigación, por lo que operó la "
             "prescripción de la acción sancionatoria. La entidad no puede sancionar conductas "
             "respecto de las cuales su potestad ha caducado."),

            ("INEXISTENCIA DE LA INFRACCIÓN",
             "Sostiene que los hechos investigados no constituyen infracción a ninguna norma "
             "legal o reglamentaria vigente. La conducta atribuida a la empresa es una práctica "
             "comercial lícita, amparada por la libertad de empresa y la autonomía privada "
             "reconocidas en la Constitución Política."),
        ]
    },
    {
        "nombre": "recurso_representante_legal.docx",
        "titulo": "RECURSO DE REPOSICIÓN — MARÍA TORRES (Representante Legal)",
        "argumentos": [
            ("NULIDAD POR VICIOS DE PROCEDIMIENTO",
             "La representante legal alega que el procedimiento administrativo adolece de "
             "vicios insubsanables de nulidad, pues el funcionario instructor no tenía "
             "competencia para firmar el pliego de cargos. Esta irregularidad afecta la "
             "validez de todos los actos expedidos con posterioridad y vicia de nulidad "
             "absoluta la resolución sancionatoria."),

            ("DESPROPORCIÓN E IRRAZONABILIDAD DE LA SANCIÓN",
             "La recurrente argumenta que la sanción impuesta es manifiestamente "
             "desproporcionada e irrazonable. La multa de cien millones de pesos equivale "
             "al 40% de los ingresos anuales de la empresa, lo que amenaza su viabilidad "
             "financiera y desconoce el principio de proporcionalidad. Solicita reducción "
             "a un máximo del 5% de los ingresos anuales."),

            ("AUSENCIA DE DOLO O CULPA GRAVE",
             "Manifiesta que la empresa actuó siempre de buena fe y no incurrió en culpa "
             "grave ni dolo. Todas las actuaciones se realizaron siguiendo el consejo de "
             "expertos legales y con la convicción de estar obrando conforme a la ley. "
             "La culpa leve no puede ser suficiente para imponer una sanción de esta "
             "magnitud en el marco de una actividad empresarial legítima."),

            ("VIOLACIÓN AL PRINCIPIO DE CONFIANZA LEGÍTIMA",
             "Señala que la Superintendencia había emitido conceptos y circulares en el "
             "pasado que permitían la conducta ahora sancionada, generando confianza "
             "legítima en los operadores del mercado. Sancionar conductas autorizadas "
             "implícitamente por la misma entidad viola los principios de buena fe y "
             "confianza legítima consagrados en la Constitución."),
        ]
    },
    {
        "nombre": "recurso_gremio_empresarial.docx",
        "titulo": "RECURSO DE REPOSICIÓN — GREMIO EMPRESARIAL (Tercero interviniente)",
        "argumentos": [
            ("IMPACTO SECTORIAL DE LA DECISIÓN",
             "El gremio empresarial manifiesta que la resolución sancionatoria crea un "
             "precedente negativo para toda la industria, pues prohíbe prácticas que han "
             "sido habituales y necesarias para el funcionamiento del mercado durante más "
             "de una década. La decisión genera inseguridad jurídica y desincentiva la "
             "inversión en el sector."),

            ("VIOLACIÓN AL DEBIDO PROCESO COLECTIVO",
             "Argumenta el gremio que no fue vinculado oportunamente al proceso, pese a "
             "tener interés legítimo en el mismo. Esta omisión desconoce su derecho de "
             "participación en actuaciones administrativas que afectan directamente a sus "
             "afiliados, conforme al artículo 13 de la Ley 1437 de 2011."),

            ("PRESCRIPCIÓN Y CADUCIDAD",
             "Reitera que la acción sancionatoria está prescrita y la entidad no podía "
             "ejercer su potestad punitiva respecto de conductas tan antiguas. La "
             "seguridad jurídica exige que las sanciones se impongan oportunamente o "
             "no se impongan. El paso del tiempo extingue la potestad sancionatoria."),

            ("FALTA DE PROPORCIONALIDAD",
             "La sanción impuesta no guarda proporcionalidad con el bien jurídico "
             "presuntamente afectado ni con las circunstancias atenuantes del caso. "
             "La ley exige que la graduación de las sanciones atienda criterios "
             "objetivos como el beneficio obtenido, el daño causado y la buena fe "
             "del infractor. Ninguno de estos criterios fue considerado adecuadamente."),
        ]
    },
]


def crear_recurso(datos):
    doc = Document()
    agregar_parrafo(doc, datos["titulo"], negrita=True, centrado=True, tamaño=13)
    agregar_parrafo(doc, "Bogotá D.C., febrero de 2024", centrado=True)
    doc.add_paragraph()

    agregar_parrafo(doc, "Señor", negrita=False)
    agregar_parrafo(doc, "SUPERINTENDENTE")
    agregar_parrafo(doc, "E. S. D.")
    doc.add_paragraph()

    agregar_parrafo(doc,
        "En ejercicio del derecho consagrado en el artículo 76 de la Ley 1437 de 2011, "
        "interpongo recurso de reposición contra la Resolución No. 2024-001, con base en "
        "los siguientes argumentos:")
    doc.add_paragraph()

    for i, (titulo, texto) in enumerate(datos["argumentos"], start=1):
        agregar_parrafo(doc, f"{i}. {titulo}", negrita=True)
        agregar_parrafo(doc, texto)
        doc.add_paragraph()

    agregar_parrafo(doc, "PETICIÓN", negrita=True)
    agregar_parrafo(doc,
        "Con base en los argumentos expuestos, solicito respetuosamente revocar la "
        "Resolución No. 2024-001 o, en subsidio, modificarla reduciendo sustancialmente "
        "la sanción impuesta.")

    ruta = os.path.join(CARPETA_RECURSOS, datos["nombre"])
    doc.save(ruta)
    print(f"Creado: {ruta}")


if __name__ == "__main__":
    crear_resolucion_base()
    for r in recursos:
        crear_recurso(r)
    print("\nDocumentos de prueba creados en:", CARPETA)
    print("  Base:", os.path.join(CARPETA_BASE, "resolucion_base.docx"))
    print("  Recursos:", CARPETA_RECURSOS)
