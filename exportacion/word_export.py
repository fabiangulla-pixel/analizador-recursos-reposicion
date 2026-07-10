"""
word_export.py
Genera un informe completo en formato Word (.docx) con todos los resultados del análisis.
Incluye: portada, reporte ejecutivo, tabla de argumentos, grupos argumentales,
propuesta de índice y registro de trazabilidad resumida.
"""

import os
from typing import List, Dict, Any
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils.logger import obtener_logger

logger = obtener_logger()

# ── Paleta de colores ────────────────────────────────────────────────────────
AZUL_OSCURO  = RGBColor(0x1A, 0x4A, 0x6E)   # encabezados principales
AZUL_MEDIO   = RGBColor(0x2C, 0x5F, 0x8A)   # encabezados secundarios
DORADO       = RGBColor(0xD4, 0xAF, 0x37)   # líneas decorativas / alertas
VERDE        = RGBColor(0x1B, 0x6B, 0x3A)   # "Sí resuelto"
NARANJA      = RGBColor(0xE6, 0x7E, 0x00)   # "Probablemente"
ROJO         = RGBColor(0xC0, 0x39, 0x2B)   # "No resuelto"
GRIS_FONDO   = "D9E2F3"                      # fondo celdas encabezado tabla (hex sin #)


# ── Utilidades de formato ────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Pone color de fondo a una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc: Document, texto: str, nivel: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    if nivel == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = AZUL_OSCURO
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif nivel == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = AZUL_MEDIO
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = AZUL_MEDIO
        run.italic = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)


def _add_paragraph(doc: Document, texto: str, italic: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(size)
    run.italic = italic
    p.paragraph_format.space_after = Pt(4)


def _color_estado(estado: str) -> RGBColor:
    e = (estado or "").upper()
    if e == "SI":
        return VERDE
    elif e == "PROBABLEMENTE":
        return NARANJA
    return ROJO


def _resumir(texto: str, max_chars: int = 250) -> str:
    if not texto:
        return ""
    if len(texto) <= max_chars:
        return texto
    return texto[:max_chars].rsplit(" ", 1)[0] + "…"


# ── Secciones del documento ──────────────────────────────────────────────────

def _portada(doc: Document, n_args: int, n_grupos: int) -> None:
    doc.add_paragraph()
    doc.add_paragraph()

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_titulo.add_run("INFORME DE ANÁLISIS\nDE RECURSOS DE REPOSICIÓN")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = AZUL_OSCURO

    doc.add_paragraph()

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_sub.add_run("Generado automáticamente por el Analizador de Recursos de Reposición")
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = AZUL_MEDIO

    doc.add_paragraph()

    p_fecha = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha = datetime.now().strftime("%d de %B de %Y — %H:%M")
    r2 = p_fecha.add_run(fecha)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # Tabla resumen rápido
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    fila = tabla.rows[0].cells
    fila[0].text = f"Total de argumentos procesados"
    fila[1].text = str(n_args)
    tabla.add_row().cells[0].text = "Grupos argumentales identificados"
    tabla.rows[1].cells[1].text = str(n_grupos)
    for row in tabla.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)

    doc.add_page_break()


def _reporte_ejecutivo(
    doc: Document,
    argumentos: List[Dict[str, Any]],
    grupos: List[Dict[str, Any]],
) -> None:
    _add_heading(doc, "I. REPORTE EJECUTIVO", 1)

    total = len(argumentos)
    resueltos   = sum(1 for a in argumentos if a.get("ya_resuelto_en_decision_base") == "SI")
    probables   = sum(1 for a in argumentos if a.get("ya_resuelto_en_decision_base") == "PROBABLEMENTE")
    nuevos      = sum(1 for a in argumentos if a.get("ya_resuelto_en_decision_base") == "NO")
    rev_humana  = sum(1 for a in argumentos if a.get("requiere_revision_humana"))
    docs        = sorted({a["archivo"] for a in argumentos})
    grp_recur   = sum(1 for g in grupos if g.get("recurrente"))

    # Tabla de métricas
    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"

    metricas = [
        ("Documentos analizados",               str(len(docs))),
        ("Total de bloques argumentativos",      str(total)),
        ("Grupos argumentales identificados",    str(len(grupos))),
        ("Grupos recurrentes (varios docs)",     str(grp_recur)),
        ("Argumentos ya resueltos en base",      str(resueltos)),
        ("Argumentos probablemente resueltos",   str(probables)),
        ("Argumentos nuevos (sin respuesta)",    str(nuevos)),
        ("Requieren revisión humana",            str(rev_humana)),
    ]

    for etiqueta, valor in metricas:
        fila = tabla.add_row().cells
        fila[0].text = etiqueta
        fila[1].text = valor
        _set_cell_bg(fila[0], GRIS_FONDO)
        for p in fila[0].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        for p in fila[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

    doc.add_paragraph()

    # Documentos
    _add_heading(doc, "Documentos analizados", 2)
    for d in docs:
        _add_paragraph(doc, f"• {d}")

    # Alertas
    _add_heading(doc, "Alertas", 2)
    alertas = []
    if nuevos > 0:
        alertas.append(f"Se identificaron {nuevos} argumentos nuevos que requieren respuesta expresa.")
    if probables > 0:
        alertas.append(f"{probables} argumentos probablemente ya resueltos — confirmar antes de omitir respuesta.")
    if rev_humana > 0:
        alertas.append(f"{rev_humana} argumentos requieren revisión humana por confianza baja o media.")
    if not alertas:
        alertas.append("No se generaron alertas.")
    for a in alertas:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(a)
        run.font.size = Pt(10)

    doc.add_page_break()


def _tabla_argumentos(doc: Document, argumentos: List[Dict[str, Any]]) -> None:
    _add_heading(doc, "II. MATRIZ DE ARGUMENTOS", 1)
    _add_paragraph(doc,
        "La siguiente tabla resume todos los argumentos extraídos de los recursos de reposición, "
        "con su grupo, estado de resolución y nivel de confianza.", italic=True)
    doc.add_paragraph()

    COLS = ["#", "Documento", "Pág.", "Grupo", "Tipo", "Texto (resumen)", "¿Resuelto?", "Confianza", "Rev. humana"]
    ANCHOS = [Cm(0.8), Cm(3.5), Cm(1.0), Cm(2.0), Cm(2.0), Cm(6.5), Cm(2.2), Cm(2.0), Cm(2.0)]

    tabla = doc.add_table(rows=1, cols=len(COLS))
    tabla.style = "Table Grid"

    # Encabezado
    hdr = tabla.rows[0].cells
    for i, (col, ancho) in enumerate(zip(COLS, ANCHOS)):
        hdr[i].text = col
        hdr[i].width = ancho
        _set_cell_bg(hdr[i], "1A4A6E")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Filas de datos
    for arg in argumentos:
        fila = tabla.add_row().cells
        estado = arg.get("ya_resuelto_en_decision_base", "")
        color_estado = _color_estado(estado)

        valores = [
            str(arg.get("id", "") + 1),
            str(arg.get("archivo", "")),
            str(arg.get("pagina", "")),
            f"G{arg.get('grupo_id', 0) + 1}",
            "Arg." if arg.get("es_argumentativo") else "Desc.",
            _resumir(arg.get("texto", ""), 200),
            estado,
            str(arg.get("confianza", "")),
            "Sí" if arg.get("requiere_revision_humana") else "No",
        ]

        for i, (celda, valor) in enumerate(zip(fila, valores)):
            celda.text = valor
            celda.width = ANCHOS[i]
            for p in celda.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
                    # Colorear la columna de estado
                    if i == 6:
                        r.font.color.rgb = color_estado
                        r.bold = True

    doc.add_page_break()


def _grupos_argumentales(doc: Document, grupos: List[Dict[str, Any]]) -> None:
    _add_heading(doc, "III. GRUPOS ARGUMENTALES", 1)
    _add_paragraph(doc,
        "Cada grupo agrupa argumentos semánticamente similares encontrados en los recursos.",
        italic=True)

    for g in grupos:
        gid = g.get("grupo_id", 0)
        n = g.get("n_argumentos", 0)
        recurrente = "Sí" if g.get("recurrente") else "No"
        estado = g.get("ya_resuelto", "NO")
        sim = g.get("similitud_base", 0.0)
        archivos = ", ".join(g.get("archivos", []))
        texto_rep = _resumir(g.get("texto_representativo", ""), 500)
        evidencia = _resumir(g.get("evidencia_base", ""), 400)

        _add_heading(doc, f"Grupo {gid + 1} — {n} argumento(s)", 2)

        tabla = doc.add_table(rows=0, cols=2)
        tabla.style = "Table Grid"

        filas_meta = [
            ("Documentos", archivos),
            ("Recurrente (varios docs)", recurrente),
            ("¿Ya resuelto en base?", estado),
            ("Similitud con base", f"{sim:.1%}"),
        ]
        for etq, val in filas_meta:
            r = tabla.add_row().cells
            r[0].text = etq
            r[1].text = val
            _set_cell_bg(r[0], GRIS_FONDO)
            for p in r[0].paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
            for p in r[1].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if etq == "¿Ya resuelto en base?":
                        run.font.color.rgb = _color_estado(val)
                        run.bold = True

        doc.add_paragraph()
        _add_heading(doc, "Argumento representativo del grupo:", 3)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(texto_rep)
        r.font.size = Pt(9)
        r.italic = True

        if evidencia:
            _add_heading(doc, "Evidencia en resolución base:", 3)
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(1)
            r2 = p2.add_run(evidencia)
            r2.font.size = Pt(9)
            r2.font.color.rgb = RGBColor(0x33, 0x66, 0x33)

        # Listado de argumentos miembro
        miembros = g.get("miembros", [])
        if miembros:
            _add_heading(doc, f"Argumentos en este grupo ({len(miembros)}):", 3)
            for m in miembros:
                p_m = doc.add_paragraph(style="List Bullet")
                p_m.paragraph_format.left_indent = Cm(0.5)
                r_m = p_m.add_run(
                    f"[{m.get('archivo','')} p.{m.get('pagina','')}] "
                    f"{_resumir(m.get('texto',''), 180)}"
                )
                r_m.font.size = Pt(8)

        doc.add_paragraph()

    doc.add_page_break()


def _propuesta_indice(doc: Document, grupos: List[Dict[str, Any]]) -> None:
    _add_heading(doc, "IV. PROPUESTA DE ÍNDICE PARA LA DECISIÓN FINAL", 1)
    _add_paragraph(doc,
        "Estructura sugerida para redactar la decisión que resuelve los recursos. "
        "Debe ser revisada y ajustada por el funcionario responsable.", italic=True)
    doc.add_paragraph()

    resueltos  = [g for g in grupos if g.get("ya_resuelto") == "SI"]
    probables  = [g for g in grupos if g.get("ya_resuelto") == "PROBABLEMENTE"]
    nuevos     = [g for g in grupos if g.get("ya_resuelto") == "NO"]

    secciones = [
        ("I. Consideraciones preliminares",
         ["Competencia para resolver", "Oportunidad de los recursos", "Legitimación de los recurrentes"]),
    ]

    _add_heading(doc, "I. Consideraciones preliminares", 2)
    for item in ["Competencia para resolver", "Oportunidad de los recursos", "Legitimación de los recurrentes"]:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item).font.size = Pt(10)

    _add_heading(doc, "II. Argumentos nuevos (requieren respuesta expresa)", 2)
    if nuevos:
        for i, g in enumerate(nuevos, 1):
            p = doc.add_paragraph(style="List Number")
            txt = f"Grupo {g['grupo_id']+1}: {_resumir(g.get('texto_representativo',''), 150)}"
            r = p.add_run(txt)
            r.font.size = Pt(10)
    else:
        _add_paragraph(doc, "(No se identificaron argumentos nuevos.)", italic=True)

    _add_heading(doc, "III. Argumentos posiblemente ya resueltos (verificar)", 2)
    if probables:
        for i, g in enumerate(probables, 1):
            p = doc.add_paragraph(style="List Number")
            txt = (f"Grupo {g['grupo_id']+1} — similitud {g.get('similitud_base',0):.1%}: "
                   f"{_resumir(g.get('texto_representativo',''), 130)}")
            p.add_run(txt).font.size = Pt(10)
    else:
        _add_paragraph(doc, "(Ninguno en esta categoría.)", italic=True)

    _add_heading(doc, "IV. Argumentos ya resueltos (confirmar y ratificar)", 2)
    if resueltos:
        for i, g in enumerate(resueltos, 1):
            p = doc.add_paragraph(style="List Number")
            txt = (f"Grupo {g['grupo_id']+1} — similitud {g.get('similitud_base',0):.1%}: "
                   f"{_resumir(g.get('texto_representativo',''), 130)}")
            p.add_run(txt).font.size = Pt(10)
    else:
        _add_paragraph(doc, "(Ninguno en esta categoría.)", italic=True)

    _add_heading(doc, "V. Decisión", 2)
    for item in ["Parte resolutiva", "Notificación"]:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item).font.size = Pt(10)

    doc.add_page_break()


def _trazabilidad_resumida(doc: Document, argumentos: List[Dict[str, Any]]) -> None:
    _add_heading(doc, "V. TRAZABILIDAD (resumen)", 1)
    _add_paragraph(doc,
        "Los siguientes registros permiten rastrear cada argumento hasta su documento y página de origen. "
        "El archivo trazabilidad.json contiene la información completa para auditoría técnica.", italic=True)
    doc.add_paragraph()

    # Solo los que requieren revisión humana
    revision = [a for a in argumentos if a.get("requiere_revision_humana")]
    if revision:
        _add_heading(doc, f"Argumentos que requieren revisión humana ({len(revision)})", 2)
        for a in revision:
            estado = a.get("ya_resuelto_en_decision_base", "")
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(
                f"[{a.get('archivo','')} p.{a.get('pagina','')}] "
                f"G{a.get('grupo_id',0)+1} | Estado: {estado} | "
                f"Similitud: {a.get('similitud_base',0):.1%} — "
                f"{_resumir(a.get('texto',''), 160)}"
            )
            r.font.size = Pt(8)
            r.font.color.rgb = NARANJA


# ── Función principal de exportación ────────────────────────────────────────

def exportar_informe_word(
    argumentos: List[Dict[str, Any]],
    grupos: List[Dict[str, Any]],
    carpeta_salida: str,
) -> None:
    """
    Genera el informe completo en Word (.docx) en la carpeta de salida.
    """
    os.makedirs(carpeta_salida, exist_ok=True)
    ruta = os.path.join(carpeta_salida, "informe_analisis.docx")

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Fuente por defecto
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    _portada(doc, len(argumentos), len(grupos))
    _reporte_ejecutivo(doc, argumentos, grupos)
    _tabla_argumentos(doc, argumentos)
    _grupos_argumentales(doc, grupos)
    _propuesta_indice(doc, grupos)
    _trazabilidad_resumida(doc, argumentos)

    doc.save(ruta)
    logger.info(f"Informe Word exportado: {ruta}")
