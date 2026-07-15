"""
borrador_decision.py
Genera un borrador REAL del acto administrativo que resuelve los recursos —
no un índice, sino el esqueleto de la decisión con los "considerandos" ya
redactados a partir de la evidencia extraída (texto del argumento, evidencia
de la resolución base, citas normativas invocadas).

Principio de honestidad (ver docs/ROADMAP.md): nada se inventa. Todo dato que
solo el funcionario conoce (número de resolución, fecha, nombre del cargo,
sentido final de la decisión) queda como marcador "[PENDIENTE: ...]" explícito,
nunca relleno plausible. Sin IA generativa: es ensamblaje de lo ya extraído.
"""

import os
from typing import Any

from docx import Document
from docx.shared import Cm, Pt

from exportacion.word_export import AZUL_OSCURO, _add_heading, _add_paragraph, _resumir
from utils.logger import obtener_logger

logger = obtener_logger()

_PENDIENTE = "[PENDIENTE: {0}]"


def _parrafo_considerando(doc: Document, numero: str, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0)
    run_num = p.add_run(f"{numero}. ")
    run_num.bold = True
    run_num.font.size = Pt(10)
    run_txt = p.add_run(texto)
    run_txt.font.size = Pt(10)


def _considerando_para_grupo(g: dict[str, Any]) -> str:
    gid = g.get("grupo_id", 0) + 1
    n = g.get("n_argumentos", 1)
    archivos = ", ".join(g.get("archivos", []))
    texto_rep = _resumir(g.get("texto_representativo", ""), 400)
    estado = g.get("ya_resuelto", "NO")
    sim = g.get("similitud_base", 0.0)
    evidencia = g.get("evidencia_base", "")

    base = f'Que en {archivos} ({n} argumento(s) en este sentido, Grupo {gid}) se plantea: "{texto_rep}". '

    if estado == "SI":
        return (
            base + "Que dicho planteamiento ya fue objeto de pronunciamiento expreso en la "
            f'resolución que se recurre, en los siguientes términos: "{_resumir(evidencia, 300)}" '
            f"(similitud {sim:.0%}). Que en consecuencia, no existe argumento nuevo que resolver "
            "y procede ratificar lo allí decidido."
        )
    if estado == "PROBABLEMENTE":
        return (
            base + "Que este planteamiento guarda similitud con lo ya considerado en la "
            f'resolución base: "{_resumir(evidencia, 300)}" (similitud {sim:.0%}), sin ser '
            f"identidad plena. {_PENDIENTE.format('el funcionario debe confirmar si el argumento ya está cubierto o requiere respuesta adicional')}."
        )
    return (
        base + f"Que dicho planteamiento no fue objeto de pronunciamiento previo. "
        f"{_PENDIENTE.format('el funcionario debe redactar la respuesta de fondo a este argumento')}."
    )


def _articulo(doc: Document, numero: str, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run_num = p.add_run(f"ARTÍCULO {numero}. ")
    run_num.bold = True
    run_num.font.size = Pt(10)
    run_txt = p.add_run(texto)
    run_txt.font.size = Pt(10)


def exportar_borrador_decision(
    argumentos: list[dict[str, Any]],
    grupos: list[dict[str, Any]],
    carpeta_salida: str,
) -> None:
    """Genera borrador_decision.docx: el esqueleto real de la decisión, no un índice."""
    os.makedirs(carpeta_salida, exist_ok=True)
    ruta = os.path.join(carpeta_salida, "borrador_decision.docx")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ── Encabezado ────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    run = p.add_run(_PENDIENTE.format("RESOLUCIÓN No. y fecha"))
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = AZUL_OSCURO

    _add_paragraph(
        doc,
        f"Por la cual se resuelven los recursos de reposición interpuestos contra la "
        f"resolución sancionatoria {_PENDIENTE.format('identificar resolución base')}",
        italic=True,
    )
    doc.add_paragraph()

    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        _PENDIENTE.format("cargo y nombre del funcionario competente")
        + ", en ejercicio de sus facultades legales y reglamentarias, y"
    )
    r2.font.size = Pt(10)
    doc.add_paragraph()

    _add_heading(doc, "CONSIDERANDO:", 1)

    docs_recurrentes = sorted({a.get("archivo", "") for a in argumentos})
    _parrafo_considerando(
        doc,
        "1",
        f"Que mediante resolución {_PENDIENTE.format('número y fecha de la resolución base')} "
        "se impuso la sanción que es objeto de los recursos que aquí se resuelven.",
    )
    _parrafo_considerando(
        doc,
        "2",
        "Que dentro del término legal se presentaron los siguientes recursos de reposición: "
        f"{', '.join(docs_recurrentes)}. "
        f"{_PENDIENTE.format('el funcionario debe confirmar competencia, oportunidad y legitimación de los recurrentes')}.",
    )

    # Argumentos nuevos primero (mayor prioridad de redacción), luego probables, luego resueltos.
    orden = {"NO": 0, "PROBABLEMENTE": 1, "SI": 2}
    grupos_ordenados = sorted(grupos, key=lambda g: orden.get(g.get("ya_resuelto", "NO"), 0))

    _parrafo_considerando(
        doc,
        "3",
        "Que analizados los argumentos presentados por los recurrentes, agrupados por "
        "similitud temática, se advierte lo siguiente:",
    )
    for i, g in enumerate(grupos_ordenados, start=1):
        _parrafo_considerando(doc, f"3.{i}", _considerando_para_grupo(g))

    citas_todas = sorted({c for a in argumentos for c in a.get("citas_normativas", [])})
    if citas_todas:
        _parrafo_considerando(
            doc,
            "4",
            "Que los recurrentes invocan las siguientes normas y providencias: "
            + "; ".join(citas_todas)
            + ". "
            + _PENDIENTE.format("el funcionario debe verificar vigencia y pertinencia de cada cita")
            + ".",
        )

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    r3 = p3.add_run("Que en mérito de lo expuesto,")
    r3.font.size = Pt(10)
    doc.add_paragraph()

    _add_heading(doc, "RESUELVE:", 1)
    _articulo(
        doc,
        "PRIMERO",
        _PENDIENTE.format(
            "confirmar, modificar o revocar la resolución recurrida, en los términos "
            "expuestos en la parte considerativa"
        )
        + ".",
    )
    _articulo(
        doc,
        "SEGUNDO",
        "Notificar la presente decisión a los recurrentes, de conformidad con lo dispuesto "
        f"en las normas aplicables. {_PENDIENTE.format('forma y datos de notificación')}.",
    )
    _articulo(
        doc,
        "TERCERO",
        _PENDIENTE.format(
            "indicar si contra la presente decisión procede algún recurso, y ante qué autoridad"
        )
        + ".",
    )

    doc.add_paragraph()
    doc.add_paragraph()
    p4 = doc.add_paragraph()
    r4 = p4.add_run(_PENDIENTE.format("NOTIFÍQUESE Y CÚMPLASE / nombre y firma del funcionario"))
    r4.bold = True
    r4.font.size = Pt(10)

    doc.save(ruta)
    n_pendientes = doc.element.body.xml.count("[PENDIENTE")
    logger.info(f"Borrador de decisión exportado: {ruta} ({n_pendientes} marcadores pendientes).")
