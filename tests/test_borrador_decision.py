"""Tests de exportacion/borrador_decision.py — esqueleto real de la decisión."""

from docx import Document

from exportacion.borrador_decision import exportar_borrador_decision


def _texto_completo(ruta: str) -> str:
    doc = Document(ruta)
    return "\n".join(p.text for p in doc.paragraphs)


def test_genera_docx_que_abre_correctamente(argumentos_procesados, grupos_procesados, tmp_path):
    exportar_borrador_decision(argumentos_procesados, grupos_procesados, str(tmp_path))
    ruta = tmp_path / "borrador_decision.docx"
    assert ruta.exists()
    doc = Document(str(ruta))  # no lanza excepción si el archivo es válido
    assert len(doc.paragraphs) > 5


def test_incluye_estructura_de_acto_administrativo(
    argumentos_procesados, grupos_procesados, tmp_path
):
    exportar_borrador_decision(argumentos_procesados, grupos_procesados, str(tmp_path))
    texto = _texto_completo(str(tmp_path / "borrador_decision.docx"))
    assert "CONSIDERANDO" in texto
    assert "RESUELVE" in texto
    assert "ARTÍCULO PRIMERO" in texto
    assert "ARTÍCULO SEGUNDO" in texto


def test_grupo_resuelto_cita_evidencia_real_no_pendiente(
    argumentos_procesados, grupos_procesados, tmp_path
):
    # grupos_procesados[0] tiene ya_resuelto="SI" con evidencia_base real (ver conftest)
    exportar_borrador_decision(argumentos_procesados, grupos_procesados, str(tmp_path))
    texto = _texto_completo(str(tmp_path / "borrador_decision.docx"))
    assert "Sobre la notificación, la resolución señaló" in texto
    assert "ratifica" in texto or "ratificar" in texto


def test_grupo_nuevo_queda_marcado_pendiente_no_inventa_respuesta(
    argumentos_procesados, grupos_procesados, tmp_path
):
    # grupos_procesados[1] tiene ya_resuelto="NO" (argumento nuevo, sin evidencia)
    exportar_borrador_decision(argumentos_procesados, grupos_procesados, str(tmp_path))
    texto = _texto_completo(str(tmp_path / "borrador_decision.docx"))
    assert "[PENDIENTE: el funcionario debe redactar la respuesta de fondo" in texto


def test_datos_desconocidos_nunca_se_inventan(argumentos_procesados, grupos_procesados, tmp_path):
    exportar_borrador_decision(argumentos_procesados, grupos_procesados, str(tmp_path))
    texto = _texto_completo(str(tmp_path / "borrador_decision.docx"))
    # Número de resolución, fecha, cargo/nombre del funcionario: nadie los conoce
    # automáticamente -> deben quedar como marcador explícito, nunca un valor plausible.
    assert "[PENDIENTE: RESOLUCIÓN No. y fecha]" in texto
    assert "[PENDIENTE: cargo y nombre del funcionario" in texto


def test_incluye_citas_normativas_agregadas(argumentos_procesados, grupos_procesados, tmp_path):
    argumentos_procesados[0]["citas_normativas"] = ["Ley 1437 de 2011, art. 76"]
    exportar_borrador_decision(argumentos_procesados, grupos_procesados, str(tmp_path))
    texto = _texto_completo(str(tmp_path / "borrador_decision.docx"))
    assert "Ley 1437 de 2011, art. 76" in texto


def test_sin_citas_no_agrega_considerando_de_normas(
    argumentos_procesados, grupos_procesados, tmp_path
):
    for a in argumentos_procesados:
        a.pop("citas_normativas", None)
    exportar_borrador_decision(argumentos_procesados, grupos_procesados, str(tmp_path))
    texto = _texto_completo(str(tmp_path / "borrador_decision.docx"))
    assert "invocan las siguientes normas" not in texto


def test_grupos_nuevos_van_antes_que_los_ya_resueltos(
    argumentos_procesados, grupos_procesados, tmp_path
):
    # El considerando 3.1 debe corresponder al grupo NO resuelto (mayor prioridad
    # de redaccion), aunque en la lista de entrada venga primero el resuelto.
    exportar_borrador_decision(argumentos_procesados, grupos_procesados, str(tmp_path))
    texto = _texto_completo(str(tmp_path / "borrador_decision.docx"))
    idx_pendiente = texto.index("respuesta de fondo")
    idx_ratifica = texto.index("ratificar lo allí decidido")
    assert idx_pendiente < idx_ratifica
